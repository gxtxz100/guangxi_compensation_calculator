#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
广西人身损害赔偿项目自动计算Web应用
基于Flask框架，提供Web界面供用户通过网络访问
"""

from flask import Flask, render_template, request, jsonify, send_file
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
import os
import tempfile
from collections import Counter

app = Flask(__name__)
app.config['SECRET_KEY'] = 'tangxuezhi'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# 2025年广西赔偿标准（根据桂高法会〔2025〕13号文件）
STANDARDS = {
    'disposable_income': 43044,  # 广西上一年度城镇居民人均可支配收入（元/年）
    'consumption': 26084,  # 广西上一年度城镇居民人均消费支出（元/年）
    'daily_meal_subsidy': 100,  # 住院伙食补助费（元/天）
    'daily_nursing_fee': 157.9,  # 护理费标准（元/天，护工标准）
    'funeral_expense': 49434,  # 丧葬费（元）
    'traffic_fee_city': 30,  # 市内交通费标准（元/天）
    'daily_accommodation_fee': 330,  # 住宿费标准（元/天）
}

# 各行业平均工资（元/年）
INDUSTRY_SALARIES = {
    '农、林、牧、渔业': 88472,
    '采矿业': 84319,
    '制造业': 81668,
    '电力、热力、燃气及水生产和供应业': 146394,
    '建筑业': 81819,
    '批发和零售业': 91322,
    '交通运输、仓储和邮政业': 116278,
    '住宿和餐饮业': 49065,
    '信息传输、软件和信息技术服务业': 140726,
    '金融业': 166109,
    '房地产业': 78846,
    '租赁和商务服务业': 74050,
    '科学研究和技术服务业': 113638,
    '水利、环境和公共设施管理业': 64797,
    '居民服务、修理和其他服务业': 56848,
    '教育': 96386,
    '卫生和社会工作': 120902,
    '文化、体育和娱乐业': 93209,
    '公共管理、社会保障和社会组织': 93976,
    '其他行业': 60000,
}

# 伤残等级系数
DISABILITY_COEFFICIENTS = {
    1: 1.0, 2: 0.9, 3: 0.8, 4: 0.7, 5: 0.6,
    6: 0.5, 7: 0.4, 8: 0.3, 9: 0.2, 10: 0.1
}


def get_float_value(value, default=0.0):
    """获取浮点数值"""
    try:
        if isinstance(value, str):
            value = value.strip()
            return float(value) if value else default
        return float(value) if value else default
    except (ValueError, TypeError):
        return default


def get_int_value(value, default=0):
    """获取整数值"""
    try:
        if isinstance(value, str):
            value = value.strip()
            return int(value) if value else default
        return int(value) if value else default
    except (ValueError, TypeError):
        return default


def calculate_compensation_years(age):
    """计算赔偿年限"""
    if age < 60:
        return 20
    elif age >= 75:
        return 5
    else:
        return 20 - (age - 60)


def calculate_multi_disability_coefficient(disability_levels_str):
    """计算多处伤残的伤残系数"""
    if not disability_levels_str or disability_levels_str.strip() == "无":
        return 1.0, None, 0.0, "无伤残，系数为1.0"
    
    disability_levels = []
    try:
        parts = disability_levels_str.replace('，', ',').replace('；', ';').replace(',', ';').split(';')
        for part in parts:
            part = part.strip()
            if not part:
                continue
            if '级' in part:
                level = int(part.replace('级', '').strip())
            else:
                level = int(part.strip())
            if 1 <= level <= 10:
                disability_levels.append(level)
    except (ValueError, AttributeError):
        return 1.0, None, 0.0, "伤残等级格式错误，按无伤残处理"
    
    if not disability_levels:
        return 1.0, None, 0.0, "无有效伤残等级，系数为1.0"
    
    level_counts = Counter(disability_levels)
    sorted_levels = sorted(level_counts.keys())
    max_level = sorted_levels[0]
    max_coefficient = DISABILITY_COEFFICIENTS.get(max_level, 1.0)
    
    if max_level == 1:
        display_levels = []
        for level, count in sorted(level_counts.items()):
            if count == 1:
                display_levels.append(f"{level}级")
            else:
                display_levels.append(f"{level}级×{count}")
        detail_parts = [f"伤残等级：{', '.join(display_levels)}\n"]
        detail_parts.append(f"最高伤残等级：1级，系数：1.00（100%）\n")
        detail_parts.append("1级伤残系数为100%，无需附加指数\n")
        detail_parts.append("最终伤残系数 = 1.00（100%）")
        detail = "".join(detail_parts)
        return 1.0, 1, 0.0, detail
    
    additional_index = 0.0
    display_levels = []
    for level, count in sorted(level_counts.items()):
        if count == 1:
            display_levels.append(f"{level}级")
        else:
            display_levels.append(f"{level}级×{count}")
    
    detail_parts = [f"伤残等级：{', '.join(display_levels)}\n"]
    detail_parts.append(f"最高伤残等级：{max_level}级，系数：{max_coefficient:.2f}\n")
    
    additional_level_info = {}
    for level in sorted_levels:
        if level == max_level:
            count = level_counts[level] - 1
            if count > 0:
                level_coefficient = DISABILITY_COEFFICIENTS.get(level, 0)
                level_additional = level_coefficient * 0.10
                total_additional = level_additional * count
                additional_index += total_additional
                additional_level_info[level] = {
                    'count': count,
                    'coefficient': level_coefficient,
                    'additional_per_unit': level_additional,
                    'total_additional': total_additional
                }
        elif level != 1:
            count = level_counts[level]
            level_coefficient = DISABILITY_COEFFICIENTS.get(level, 0)
            level_additional = level_coefficient * 0.10
            total_additional = level_additional * count
            additional_index += total_additional
            additional_level_info[level] = {
                'count': count,
                'coefficient': level_coefficient,
                'additional_per_unit': level_additional,
                'total_additional': total_additional
            }
    
    if additional_level_info:
        detail_parts.append("附加伤残等级：")
        info_list = []
        for level in sorted(additional_level_info.keys()):
            info = additional_level_info[level]
            if info['count'] == 1:
                info_list.append(f"{level}级（赔偿系数{info['coefficient']:.2f}，附加{info['additional_per_unit']*100:.2f}%）")
            else:
                info_list.append(f"{level}级×{info['count']}（赔偿系数{info['coefficient']:.2f}，每处附加{info['additional_per_unit']*100:.2f}%，合计{info['total_additional']*100:.2f}%）")
        detail_parts.append("、".join(info_list))
        
        original_additional_index = additional_index
        additional_index = min(additional_index, 0.10)
        
        if original_additional_index > 0.10:
            detail_parts.append(f"\n附加指数合计：{original_additional_index * 100:.2f}%，超过10%上限，按10%计算\n")
        else:
            detail_parts.append(f"\n附加指数合计：{additional_index * 100:.2f}%\n")
    else:
        detail_parts.append("无附加伤残等级\n")
    
    final_coefficient = min(max_coefficient + additional_index, 1.0)
    detail_parts.append(f"最终伤残系数 = {max_coefficient:.2f} + {additional_index:.2f} = {final_coefficient:.2f}")
    if final_coefficient >= 1.0:
        detail_parts.append("（已达到100%上限）")
    
    detail = "".join(detail_parts)
    return final_coefficient, max_level, additional_index, detail


def calculate_work_loss_fee(data):
    """计算误工费"""
    work_loss_days = get_int_value(data.get('work_loss_days', 0))
    if work_loss_days <= 0:
        return 0, "误工天数为0，不计算误工费"
    
    income_type = data.get('work_income_type', '固定收入')
    
    if income_type == "固定收入":
        monthly_income = get_float_value(data.get('monthly_income', 0))
        if monthly_income > 0:
            daily_income = monthly_income / 30
            amount = daily_income * work_loss_days
            detail = f"固定收入计算：\n月收入：{monthly_income:,.2f}元\n日均收入 = 月收入 ÷ 30 = {monthly_income:,.2f} ÷ 30 = {daily_income:,.2f}元/天\n误工费 = 日均收入 × 误工天数 = {daily_income:,.2f} × {work_loss_days} = {amount:,.2f}元"
            return amount, detail
        else:
            return 0, "月收入为0，不计算误工费"
    
    elif income_type == "无固定收入（能证明最近三年平均）":
        avg_daily_income = get_float_value(data.get('avg_daily_income', 0))
        if avg_daily_income > 0:
            amount = avg_daily_income * work_loss_days
            detail = f"无固定收入（能证明最近三年平均）计算：\n最近三年平均日均收入：{avg_daily_income:,.2f}元/天\n误工费 = 日均收入 × 误工天数 = {avg_daily_income:,.2f} × {work_loss_days} = {amount:,.2f}元"
            return amount, detail
        else:
            return 0, "日均收入为0，不计算误工费"
    
    else:
        selected_industry = data.get('industry_type', '其他行业')
        industry_avg_salary = INDUSTRY_SALARIES.get(selected_industry, INDUSTRY_SALARIES['其他行业'])
        daily_avg_salary = industry_avg_salary / 365
        amount = daily_avg_salary * work_loss_days
        detail = f"无固定收入（不能证明，参照行业平均）计算\n选择行业：{selected_industry}\n行业平均工资：{industry_avg_salary:,.2f}元/年\n日均工资 = 年工资 ÷ 365 = {industry_avg_salary:,.2f} ÷ 365 = {daily_avg_salary:,.2f}元/天\n误工费 = 日均工资 × 误工天数 = {daily_avg_salary:,.2f} × {work_loss_days} = {amount:,.2f}元"
        return amount, detail


def calculate_nursing_fee(data):
    """计算护理费"""
    nursing_days = get_int_value(data.get('nursing_days', 0))
    nursing_count = get_int_value(data.get('nursing_count', 1))
    
    if nursing_days <= 0:
        return 0, "护理天数为0，不计算护理费"
    
    nursing_type = data.get('nursing_type', '无收入或雇佣护工')
    
    if nursing_type == "有收入":
        nursing_income = get_float_value(data.get('nursing_income', 0))
        if nursing_income > 0:
            amount = nursing_income * nursing_days * nursing_count
            detail = f"护理人员有收入计算：\n护理人员日均收入：{nursing_income:,.2f}元/天\n护理天数：{nursing_days}天\n护理人数：{nursing_count}人\n护理费 = 日均收入 × 护理天数 × 护理人数 = {nursing_income:,.2f} × {nursing_days} × {nursing_count} = {amount:,.2f}元"
            return amount, detail
        else:
            return 0, "护理人员日均收入为0，不计算护理费"
    else:
        nursing_fee_per_day = STANDARDS['daily_nursing_fee']
        amount = nursing_fee_per_day * nursing_days * nursing_count
        detail = f"无收入或雇佣护工计算：\n护工标准：{nursing_fee_per_day:,.2f}元/天\n护理天数：{nursing_days}天\n护理人数：{nursing_count}人\n护理费 = 护工标准 × 护理天数 × 护理人数 = {nursing_fee_per_day:,.2f} × {nursing_days} × {nursing_count} = {amount:,.2f}元"
        return amount, detail


def calculate_dependent_living_expense(data, victim_age, disability_coefficient=1.0, is_death=False):
    """计算被扶养人生活费"""
    dependent_info_str = data.get('dependent_info', '').strip()
    if not dependent_info_str:
        return 0, "未填写被扶养人信息，不计算被扶养人生活费"
    
    base_consumption = STANDARDS['consumption']
    consumption_type = "广西上一年度城镇居民人均消费支出"
    
    dependents = []
    try:
        for item in dependent_info_str.split(';'):
            item = item.strip()
            if not item:
                continue
            if ',' in item:
                parts = item.split(',')
                age = int(parts[0].strip())
                support_count = int(parts[1].strip()) if len(parts) > 1 else 1
                dependents.append({'age': age, 'support_count': support_count})
            else:
                age = int(item)
                dependents.append({'age': age, 'support_count': 1})
    except ValueError:
        return 0, "被扶养人信息格式错误"
    
    if not dependents:
        return 0, "未填写被扶养人信息，不计算被扶养人生活费"
    
    dependent_expenses = []
    detail_parts = [f"{consumption_type}：{base_consumption:,.2f}元/年\n"]
    
    for idx, dep in enumerate(dependents):
        age = dep['age']
        support_count = dep['support_count']
        
        if age < 18:
            years = 18 - age
            age_desc = f"不满18周岁，按(18-{age})年计算"
        elif age >= 18 and age < 60:
            years = 20
            age_desc = f"18-60周岁（无劳动能力），按20年计算"
        elif age >= 60 and age < 75:
            years = 20 - (age - 60)
            age_desc = f"60-75周岁，按[20-({age}-60)]={years}年计算"
        else:
            years = 5
            age_desc = f"75周岁以上，按5年计算"
        
        if years <= 0:
            continue
        
        annual_expense_per_dependent = base_consumption / support_count
        dependent_expenses.append({
            'age': age,
            'years': years,
            'support_count': support_count,
            'annual_expense': annual_expense_per_dependent
        })
        
        detail_parts.append(f"被扶养人{idx+1}：{age}岁，{age_desc}，扶养人数{support_count}人\n年生活费 = {base_consumption:,.2f} ÷ {support_count} = {annual_expense_per_dependent:,.2f}元/年\n")
    
    if not dependent_expenses:
        return 0, "被扶养人信息无效"
    
    max_years = max(exp['years'] for exp in dependent_expenses)
    total_expense = 0
    year_details = []
    
    for year in range(max_years):
        year_total = 0
        active_deps = []
        for exp in dependent_expenses:
            if year < exp['years']:
                year_total += exp['annual_expense']
                active_deps.append(f"{exp['age']}岁")
        
        original_total = year_total
        year_total = min(year_total, base_consumption)
        total_expense += year_total
        
        if year_total > 0:
            if original_total > base_consumption:
                year_details.append(f"第{year+1}年：{'+'.join(active_deps)}的年生活费合计{original_total:,.2f}元，超过{base_consumption:,.2f}元，按{base_consumption:,.2f}元计算")
            else:
                year_details.append(f"第{year+1}年：{'+'.join(active_deps)}的年生活费合计{year_total:,.2f}元")
    
    year_amounts = []
    for year in range(max_years):
        year_total = 0
        for exp in dependent_expenses:
            if year < exp['years']:
                year_total += exp['annual_expense']
        year_total = min(year_total, base_consumption)
        if year_total > 0:
            year_amounts.append(f"{year_total:,.2f}")
    
    total_formula = " + ".join(year_amounts) if year_amounts else "0"
    original_total = total_expense
    total_expense = total_expense * disability_coefficient
    
    if is_death:
        detail = "".join(detail_parts) + "\n按年计算明细：\n" + "\n".join(year_details) + f"\n\n小计 = " + total_formula + f" = {original_total:,.2f}元\n受害人死亡，系数为100%（无需乘以伤残系数）\n被扶养人生活费 = 小计 × 100% = {original_total:,.2f} × 1.0 = {total_expense:,.2f}元"
    elif disability_coefficient < 1.0:
        detail = "".join(detail_parts) + "\n按年计算明细：\n" + "\n".join(year_details) + f"\n\n小计 = " + total_formula + f" = {original_total:,.2f}元\n伤残系数：{disability_coefficient:.2f}\n被扶养人生活费 = 小计 × 伤残系数 = {original_total:,.2f} × {disability_coefficient:.2f} = {total_expense:,.2f}元"
    else:
        detail = "".join(detail_parts) + "\n按年计算明细：\n" + "\n".join(year_details) + f"\n\n总计 = " + total_formula + f" = {total_expense:,.2f}元"
    
    return total_expense, detail


@app.route('/')
def index():
    """主页"""
    return render_template('index.html', 
                         industry_salaries=list(INDUSTRY_SALARIES.keys()),
                         standards=STANDARDS,
                         current_date=datetime.now().strftime('%Y-%m-%d'))


@app.route('/api/calculate', methods=['POST'])
def calculate():
    """计算赔偿API"""
    try:
        data = request.json
        results = {}
        calculation_details = {}
        
        victim_name = data.get('victim_name', '').strip() or "未填写"
        victim_age = get_int_value(data.get('victim_age', 0))
        
        # 1. 医疗费
        medical_expense = get_float_value(data.get('medical_expense', 0))
        results['医疗费'] = medical_expense
        if medical_expense > 0:
            calculation_details['医疗费'] = f"医疗费 = 诊疗费 + 医药费 + 住院费 = {medical_expense:,.2f}元"
        
        # 2. 后续治疗费
        follow_up_treatment_fee = get_float_value(data.get('follow_up_treatment_fee', 0))
        results['后续治疗费'] = follow_up_treatment_fee
        if follow_up_treatment_fee > 0:
            calculation_details['后续治疗费'] = f"后续治疗费 = {follow_up_treatment_fee:,.2f}元"
        
        # 3. 住院伙食补助费
        hospital_days = get_int_value(data.get('hospital_days', 0))
        meal_subsidy_per_day = get_float_value(data.get('meal_subsidy', STANDARDS['daily_meal_subsidy']))
        meal_subsidy_total = hospital_days * meal_subsidy_per_day
        results['住院伙食补助费'] = meal_subsidy_total
        if meal_subsidy_total > 0:
            calculation_details['住院伙食补助费'] = f"住院天数：{hospital_days}天\n补助标准：{meal_subsidy_per_day:,.2f}元/天\n住院伙食补助费 = 住院天数 × 补助标准 = {hospital_days} × {meal_subsidy_per_day:,.2f} = {meal_subsidy_total:,.2f}元"
        
        # 4. 营养费
        nutrition_fee = get_float_value(data.get('nutrition_fee', 0))
        results['营养费'] = nutrition_fee
        if nutrition_fee > 0:
            calculation_details['营养费'] = f"营养费 = {nutrition_fee:,.2f}元"
        
        # 5. 交通费
        traffic_fee = get_float_value(data.get('traffic_fee', 0))
        results['交通费'] = traffic_fee
        if traffic_fee > 0:
            calculation_details['交通费'] = f"交通费 = {traffic_fee:,.2f}元"
        
        # 6. 住宿费
        accommodation_days = get_int_value(data.get('accommodation_days', 0))
        accommodation_fee_per_day = STANDARDS['daily_accommodation_fee']
        accommodation_fee = accommodation_days * accommodation_fee_per_day
        results['住宿费'] = accommodation_fee
        if accommodation_fee > 0:
            calculation_details['住宿费'] = f"住宿天数：{accommodation_days}天\n住宿费标准：{accommodation_fee_per_day:,.2f}元/天\n住宿费 = 住宿天数 × 住宿费标准 = {accommodation_days} × {accommodation_fee_per_day:,.2f} = {accommodation_fee:,.2f}元"
        
        # 7. 误工费
        work_loss_fee, work_detail = calculate_work_loss_fee(data)
        results['误工费'] = work_loss_fee
        calculation_details['误工费'] = work_detail
        
        # 8. 护理费
        nursing_fee_total, nursing_detail = calculate_nursing_fee(data)
        results['护理费'] = nursing_fee_total
        calculation_details['护理费'] = nursing_detail
        
        # 9. 残疾赔偿金
        disability_level_str = data.get('disability_level', '').strip() or "无"
        disability_coefficient, max_level, additional_index, disability_detail = \
            calculate_multi_disability_coefficient(disability_level_str)
        
        if disability_coefficient < 1.0 or (disability_level_str and disability_level_str != "无"):
            base_income = STANDARDS['disposable_income']
            income_type = "广西上一年度城镇居民人均可支配收入"
            years = calculate_compensation_years(victim_age)
            disability_compensation = base_income * years * disability_coefficient
            results['残疾赔偿金'] = disability_compensation
            year_desc = f"{years}年" if victim_age < 60 else (f"{years}年（60周岁以上每增加一岁减少一年）" if victim_age < 75 else f"{years}年（75周岁以上按5年计算）")
            detail = f"{disability_detail}\n{income_type}：{base_income:,.2f}元/年\n赔偿年限：{year_desc}\n残疾赔偿金 = {income_type} × 赔偿年限 × 伤残系数 = {base_income:,.2f} × {years} × {disability_coefficient:.2f} = {disability_compensation:,.2f}元"
            calculation_details['残疾赔偿金'] = detail
        else:
            results['残疾赔偿金'] = 0
        
        # 10. 残疾辅助器具费
        disability_appliance_fee = get_float_value(data.get('disability_appliance_fee', 0))
        results['残疾辅助器具费'] = disability_appliance_fee
        if disability_appliance_fee > 0:
            calculation_details['残疾辅助器具费'] = f"残疾辅助器具费 = {disability_appliance_fee:,.2f}元"
        
        # 11. 被扶养人生活费
        is_death = data.get('is_death', False)
        if is_death:
            dependent_coefficient = 1.0
        else:
            dependent_coefficient = disability_coefficient
        
        dependent_living_expense, dependent_detail = calculate_dependent_living_expense(
            data, victim_age, dependent_coefficient, is_death)
        results['被扶养人生活费'] = dependent_living_expense
        if dependent_living_expense > 0:
            calculation_details['被扶养人生活费'] = dependent_detail
        
        # 12. 死亡赔偿金
        if is_death:
            base_income = STANDARDS['disposable_income']
            income_type = "广西上一年度城镇居民人均可支配收入"
            years = calculate_compensation_years(victim_age)
            death_compensation = base_income * years
            results['死亡赔偿金'] = death_compensation
            results['丧葬费'] = STANDARDS['funeral_expense']
            year_desc = f"{years}年" if victim_age < 60 else (f"{years}年（60周岁以上每增加一岁减少一年）" if victim_age < 75 else f"{years}年（75周岁以上按5年计算）")
            calculation_details['死亡赔偿金'] = f"{income_type}：{base_income:,.2f}元/年\n赔偿年限：{year_desc}\n死亡赔偿金 = {income_type} × 赔偿年限 = {base_income:,.2f} × {years} = {death_compensation:,.2f}元"
            calculation_details['丧葬费'] = f"丧葬费 = {STANDARDS['funeral_expense']:,.2f}元"
        else:
            results['死亡赔偿金'] = 0
            results['丧葬费'] = 0
        
        # 13. 精神损害抚慰金
        mental_damage = get_float_value(data.get('mental_damage', 0))
        results['精神损害抚慰金'] = mental_damage
        if mental_damage > 0:
            calculation_details['精神损害抚慰金'] = f"精神损害抚慰金 = {mental_damage:,.2f}元"
        
        # 计算总计
        total = sum(results.values())
        results['总计'] = total
        
        valid_items = [item for item in ['医疗费', '后续治疗费', '误工费', '护理费', '交通费', '住宿费', '住院伙食补助费', 
                      '营养费', '残疾赔偿金', '残疾辅助器具费', '被扶养人生活费', 
                      '死亡赔偿金', '丧葬费', '精神损害抚慰金']
                      if item in results and results[item] > 0]
        total_formula = " + ".join([f"{results[item]:,.2f}" for item in valid_items])
        calculation_details['总计'] = f"总计 = {total_formula} = {total:,.2f}元"
        
        return jsonify({
            'success': True,
            'results': results,
            'details': calculation_details,
            'victim_name': victim_name,
            'victim_age': victim_age
        })
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/export_word', methods=['POST'])
def export_word():
    """导出Word文档API"""
    try:
        data = request.json
        results = data.get('results', {})
        details = data.get('details', {})
        victim_name = data.get('victim_name', '未填写')
        victim_age = data.get('victim_age', 0)
        accident_date = data.get('accident_date', datetime.now().strftime('%Y-%m-%d'))
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx', dir='/app/temp')
        temp_file.close()
        filename = temp_file.name
        
        # 创建Word文档
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 设置标题样式
        heading1 = doc.styles['Heading 1']
        heading1_font = heading1.font
        heading1_font.name = '黑体'
        heading1_font.size = Pt(16)
        heading1_font.bold = True
        heading1_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        heading2 = doc.styles['Heading 2']
        heading2_font = heading2.font
        heading2_font.name = '黑体'
        heading2_font.size = Pt(14)
        heading2_font.bold = True
        heading2_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加页脚（页码）
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = footer_para._element
        p_r = OxmlElement('w:pPr')
        p.append(p_r)
        
        r = OxmlElement('w:r')
        p.append(r)
        
        t = OxmlElement('w:t')
        t.text = '第 '
        r.append(t)
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r.append(fldChar2)
        
        r2 = OxmlElement('w:r')
        p.append(r2)
        t2 = OxmlElement('w:t')
        t2.text = ' 页'
        r2.append(t2)
        
        for r_elem in p.findall(qn('w:r')):
            rPr = OxmlElement('w:rPr')
            r_elem.insert(0, rPr)
            font = OxmlElement('w:rFonts')
            font.set(qn('w:ascii'), '宋体')
            font.set(qn('w:eastAsia'), '宋体')
            font.set(qn('w:hAnsi'), '宋体')
            rPr.append(font)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')
            rPr.append(sz)
        
        # 标题
        title = doc.add_heading('广西人身损害赔偿计算结果', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = '黑体'
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        doc.add_paragraph()
        
        # 基本信息
        doc.add_heading('一、基本信息', level=1)
        basic_table = doc.add_table(rows=3, cols=2)
        basic_table.style = 'Light Grid Accent 1'
        basic_table.columns[0].width = Inches(2.0)
        basic_table.columns[1].width = Inches(4.5)
        
        basic_info = [
            ('受害人姓名', victim_name),
            ('受害人年龄', f"{victim_age}岁"),
            ('事故发生日期', accident_date),
        ]
        
        for i, (label, value) in enumerate(basic_info):
            label_cell = basic_table.rows[i].cells[0]
            label_cell.text = label
            label_para = label_cell.paragraphs[0]
            label_para.runs[0].bold = True
            label_para.runs[0].font.name = '宋体'
            label_para.runs[0].font.size = Pt(12)
            label_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            value_cell = basic_table.rows[i].cells[1]
            value_cell.text = value
            value_para = value_cell.paragraphs[0]
            value_para.runs[0].font.name = '宋体'
            value_para.runs[0].font.size = Pt(12)
            value_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        doc.add_paragraph()
        
        # 赔偿明细
        doc.add_heading('二、赔偿明细及计算公式', level=1)
        items_order = ['医疗费', '后续治疗费', '误工费', '护理费', '交通费', '住宿费', '住院伙食补助费', 
                      '营养费', '残疾赔偿金', '残疾辅助器具费', '被扶养人生活费', 
                      '死亡赔偿金', '丧葬费', '精神损害抚慰金']
        
        valid_items = [item for item in items_order if item in results and results[item] > 0]
        
        if valid_items:
            detail_table = doc.add_table(rows=len(valid_items) + 1, cols=4)
            detail_table.style = 'Light Grid Accent 1'
            
            tbl = detail_table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '0')
            tblW.set(qn('w:type'), 'auto')
            tblPr.append(tblW)
            
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'autofit')
            tblPr.append(tblLayout)
            
            detail_table.columns[0].width = Inches(0.4)
            detail_table.columns[1].width = Inches(1.0)
            detail_table.columns[2].width = Inches(1.0)
            detail_table.columns[3].width = Inches(5.1)
            
            for row_idx, row in enumerate(detail_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    tcPr = cell._element.tcPr
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        cell._element.insert(0, tcPr)
                    
                    if col_idx < 3:
                        left_right = '80'
                        top_bottom = '50'
                    else:
                        left_right = '120'
                        top_bottom = '60'
                    
                    tcMar = OxmlElement('w:tcMar')
                    for margin_name, margin_value in [('top', top_bottom), ('left', left_right), 
                                                      ('bottom', top_bottom), ('right', left_right)]:
                        margin = OxmlElement(f'w:{margin_name}')
                        margin.set(qn('w:w'), margin_value)
                        margin.set(qn('w:type'), 'dxa')
                        tcMar.append(margin)
                    tcPr.append(tcMar)
            
            header_cells = detail_table.rows[0].cells
            header_texts = ['序号', '项目', '金额', '计算方式']
            
            for idx, cell in enumerate(header_cells):
                header_text = header_texts[idx]
                cell.paragraphs[0].clear()
                para = cell.paragraphs[0]
                run = para.add_run(header_text)
                run.bold = True
                run.font.name = '黑体'
                run.font.size = Pt(11)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                tcPr = cell._element.get_or_add_tcPr()
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'E7E6E6')
                shading_elm.set(qn('w:val'), 'clear')
                tcPr.append(shading_elm)
            
            for idx, item in enumerate(valid_items):
                row = detail_table.rows[idx + 1]
                
                cell0 = row.cells[0]
                cell0.paragraphs[0].clear()
                para0 = cell0.paragraphs[0]
                run0 = para0.add_run(str(idx + 1))
                run0.font.name = '宋体'
                run0.font.size = Pt(10)
                run0._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                cell1 = row.cells[1]
                cell1.paragraphs[0].clear()
                para1 = cell1.paragraphs[0]
                run1 = para1.add_run(item)
                run1.font.name = '宋体'
                run1.font.size = Pt(10)
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                cell2 = row.cells[2]
                cell2.paragraphs[0].clear()
                para2 = cell2.paragraphs[0]
                amount_text = f"{results[item]:,.2f}"
                run2 = para2.add_run(amount_text)
                run2.font.name = '宋体'
                run2.font.size = Pt(10)
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                cell3 = row.cells[3]
                cell3.paragraphs[0].clear()
                para3 = cell3.paragraphs[0]
                
                if item in details:
                    detail = details[item]
                    if '\n' in detail:
                        lines = detail.split('\n')
                        for i, line in enumerate(lines):
                            if i > 0:
                                para3 = cell3.add_paragraph()
                            run3 = para3.add_run(line.strip())
                            run3.font.name = '宋体'
                            run3.font.size = Pt(9.5)
                            run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        formula_text = detail.replace('；', '\n').replace(';', '\n')
                        if '\n' in formula_text:
                            lines = formula_text.split('\n')
                            for i, line in enumerate(lines):
                                if i > 0:
                                    para3 = cell3.add_paragraph()
                                run3 = para3.add_run(line.strip())
                                run3.font.name = '宋体'
                                run3.font.size = Pt(9.5)
                                run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        else:
                            run3 = para3.add_run(formula_text)
                            run3.font.name = '宋体'
                            run3.font.size = Pt(9.5)
                            run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    run3 = para3.add_run(f"{item} = {results[item]:,.2f} 元")
                    run3.font.name = '宋体'
                    run3.font.size = Pt(9.5)
                    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                para3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                tr = row._element
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), '300')
                trHeight.set(qn('w:hRule'), 'atLeast')
                trPr.append(trHeight)
            
            doc.add_paragraph()
        
        # 总计
        doc.add_heading('三、赔偿总额', level=1)
        total_table = doc.add_table(rows=2, cols=2)
        total_table.style = 'Light Grid Accent 1'
        total_table.columns[0].width = Inches(2.0)
        total_table.columns[1].width = Inches(5.0)
        
        total_table.rows[0].cells[0].text = '项目'
        total_table.rows[0].cells[1].text = '金额（元）'
        for cell in total_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        total_table.rows[1].cells[0].text = '赔偿总额'
        total_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
        total_table.rows[1].cells[1].text = f"{results.get('总计', 0):,.2f}"
        total_table.rows[1].cells[1].paragraphs[0].runs[0].bold = True
        total_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if '总计' in details:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('计算公式：').bold = True
            doc.add_paragraph(details['总计'])
        
        # 计算依据
        doc.add_heading('四、计算依据', level=1)
        doc.add_paragraph('本计算依据以下法律法规及标准文件：')
        doc.add_paragraph('《广西壮族自治区道路交通事故损害赔偿项目及计算标准》（桂高法会〔2025〕13号）', style='List Number')
        doc.add_paragraph('《广西壮族自治区公安厅关于道路交通事故处理有关问题的通知》（桂公通〔2025〕60号）', style='List Number')
        doc.add_paragraph()
        doc.add_paragraph('注：2025年标准统一使用广西上一年度城镇居民人均可支配收入和城镇居民人均消费支出标准进行计算。')
        
        # 备注
        doc.add_heading('五、备注', level=1)
        doc.add_paragraph('1. 本计算结果仅供参考，实际赔偿金额以法院判决为准。')
        doc.add_paragraph('2. 各项费用需提供相应的票据和证明材料。')
        doc.add_paragraph('3. 误工费、护理费的计算方式已根据收入类型进行区分。')
        doc.add_paragraph('4. 被扶养人生活费的计算已考虑年赔偿总额限制。')
        doc.add_paragraph('5. 如对计算结果有疑问，请咨询广西瀛桂律师事务所唐学智律师，联系电话18078374299。')
        
        # 保存文档
        doc.save(filename)
        
        return send_file(filename, as_attachment=True, 
                       download_name=f"{victim_name if victim_name != '未填写' else '赔偿'}计算结果.docx",
                       mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)

