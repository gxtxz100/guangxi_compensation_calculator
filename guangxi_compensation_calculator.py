#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
广西人身损害赔偿项目自动计算程序
根据《最高人民法院关于审理人身损害赔偿案件适用法律若干问题的解释》及相关标准计算各项赔偿项目并生成Word文档
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
import os


class GuangxiCompensationCalculator:
    """广西人身损害赔偿计算器"""
    
    # 2025年广西赔偿标准（根据桂高法会〔2025〕13号文件）
    # 注意：2025年标准统一使用城镇居民人均可支配收入，不再区分城镇和农村
    STANDARDS = {
        'disposable_income': 43044,  # 广西上一年度城镇居民人均可支配收入（元/年）
        'consumption': 26084,  # 广西上一年度城镇居民人均消费支出（元/年）
        'daily_meal_subsidy': 100,  # 住院伙食补助费（元/天）
        'daily_nursing_fee': 157.9,  # 护理费标准（元/天，护工标准）
        'funeral_expense': 49434,  # 丧葬费（元）
        'traffic_fee_city': 30,  # 市内交通费标准（元/天）
        'daily_accommodation_fee': 330,  # 住宿费标准（元/天）
    }
    
    # 各行业平均工资（元/年）- 根据广西上一年度统计数据
    # 数据来源：桂公通〔2025〕60号文件
    INDUSTRY_SALARIES = {
        '农、林、牧、渔业': 88472,
        '采矿业': 84319,
        '制造业': 81668,
        '电力、热力、燃气及水生产和供应业': 146394,
        '建筑业': 81819,
        '批发和零售业': 91322,
        '交通运输、仓储和邮政业':116278,
        '住宿和餐饮业': 49065,
        '信息传输、软件和信息技术服务业': 140726,
        '金融业': 166109,
        '房地产业': 78846,
        '租赁和商务服务业': 74050,
        '科学研究和技术服务业': 113638,
        '水利、环境和公共设施管理业': 64797,
        '居民服务、修理和其他服务业': 56848,
        '教育': 96386,
        '卫生和社会工作': 120902,
        '文化、体育和娱乐业': 93209,
        '公共管理、社会保障和社会组织': 93976,
        '其他行业': 60000,
    }
    
    # 伤残等级系数
    DISABILITY_COEFFICIENTS = {
        1: 1.0,
        2: 0.9,
        3: 0.8,
        4: 0.7,
        5: 0.6,
        6: 0.5,
        7: 0.4,
        8: 0.3,
        9: 0.2,
        10: 0.1
    }
    
    def __init__(self, root):
        self.root = root
        self.root.title("广西瀛桂律师事务所 唐学智律师制作 18078374299")
        self.root.geometry("900x1000")
        self.root.resizable(True, True)
        
        # 创建主框架
        self.create_widgets()
        
    def create_widgets(self):
        """创建GUI组件"""
        # 创建滚动框架
        canvas = tk.Canvas(self.root)
        scrollbar = ttk.Scrollbar(self.root, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        # 更新canvas的scrollregion
        def update_scrollregion(event=None):
            canvas.update_idletasks()
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        scrollable_frame.bind("<Configure>", update_scrollregion)
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 绑定鼠标滚轮事件 - Mac特殊处理
        import platform
        system = platform.system()
        
        def _on_mousewheel(event):
            """处理鼠标滚轮事件"""
            try:
                if system == "Darwin":  # Mac系统
                    # Mac使用delta，值通常是正数向上，负数向下
                    # 但需要除以一个系数来调整滚动速度
                    delta = getattr(event, 'delta', 0)
                    if delta:
                        # Mac的delta值较小，需要调整
                        scroll_amount = int(-1 * delta / 3)  # 调整滚动速度
                        canvas.yview_scroll(scroll_amount, "units")
                elif system == "Windows":  # Windows系统
                    # Windows使用delta/120
                    delta = getattr(event, 'delta', 0)
                    if delta:
                        canvas.yview_scroll(int(-1 * (delta / 120)), "units")
                else:  # Linux系统
                    # Linux使用Button-4和Button-5
                    num = getattr(event, 'num', 0)
                    if num == 4:
                        canvas.yview_scroll(-1, "units")
                    elif num == 5:
                        canvas.yview_scroll(1, "units")
            except Exception as e:
                # 如果出错，尝试通用方法
                try:
                    delta = getattr(event, 'delta', 0)
                    if delta:
                        canvas.yview_scroll(int(-1 * delta / 3), "units")
                except:
                    pass
        
        # Mac系统需要特殊处理 - 直接绑定，不需要Enter/Leave事件
        if system == "Darwin":
            # Mac上直接绑定到所有组件，不需要Enter/Leave
            # 使用bind_all确保全局绑定
            self.root.bind_all("<MouseWheel>", _on_mousewheel)
            self.root.bind_all("<Button-4>", _on_mousewheel)
            self.root.bind_all("<Button-5>", _on_mousewheel)
            
            # 同时也绑定到canvas本身
            canvas.bind("<MouseWheel>", _on_mousewheel)
            canvas.bind("<Button-4>", _on_mousewheel)
            canvas.bind("<Button-5>", _on_mousewheel)
            
            # 绑定到scrollable_frame
            scrollable_frame.bind("<MouseWheel>", _on_mousewheel)
            scrollable_frame.bind("<Button-4>", _on_mousewheel)
            scrollable_frame.bind("<Button-5>", _on_mousewheel)
        else:
            # Windows和Linux的处理
            def _bind_to_mousewheel(event):
                if system == "Windows":
                    canvas.bind_all("<MouseWheel>", _on_mousewheel)
                else:  # Linux
                    canvas.bind_all("<Button-4>", _on_mousewheel)
                    canvas.bind_all("<Button-5>", _on_mousewheel)
            
            def _unbind_from_mousewheel(event):
                canvas.unbind_all("<MouseWheel>")
                canvas.unbind_all("<Button-4>")
                canvas.unbind_all("<Button-5>")
            
            canvas.bind('<Enter>', _bind_to_mousewheel)
            canvas.bind('<Leave>', _unbind_from_mousewheel)
            scrollable_frame.bind('<Enter>', _bind_to_mousewheel)
            scrollable_frame.bind('<Leave>', _unbind_from_mousewheel)
            self.root.bind('<Enter>', _bind_to_mousewheel)
        
        # 确保canvas可以接收焦点
        canvas.focus_set()
        
        # 保存canvas引用以便后续使用
        self.canvas = canvas
        self.scrollable_frame = scrollable_frame
        
        # 标题区域 - 紧凑设计
        title_frame = tk.Frame(scrollable_frame, bg="#2c3e50", height=45)
        title_frame.pack(fill="x", padx=0, pady=0)
        title_label = tk.Label(title_frame, text="广西人身损害赔偿计算器", 
                               font=("Microsoft YaHei", 16, "bold"),
                               bg="#2c3e50", fg="white")
        title_label.pack(pady=8)
        
        # 副标题
        subtitle_label = tk.Label(title_frame, 
                                 text="根据（桂高法会〔2025〕13号），（桂公通〔2025〕60号）",
                                 font=("Microsoft YaHei", 8),
                                 bg="#2c3e50", fg="#ecf0f1")
        subtitle_label.pack(pady=(0, 5))
        
        # 基本信息框架 - 紧凑设计
        basic_frame = ttk.LabelFrame(scrollable_frame, text="📋 基本信息", padding=6)
        basic_frame.pack(fill="x", padx=10, pady=3)
        
        self.victim_name = self.create_entry(basic_frame, "受害人姓名：", 0)
        self.victim_age = self.create_entry(basic_frame, "受害人年龄：", 1)
        self.accident_date_year, self.accident_date_month, self.accident_date_day = \
            self.create_date_selectors(basic_frame, "事故发生日期：", 2)
        
        # 医疗相关费用框架
        medical_frame = ttk.LabelFrame(scrollable_frame, text="🏥 医疗相关费用", padding=6)
        medical_frame.pack(fill="x", padx=10, pady=3)
        
        self.medical_expense = self.create_entry(medical_frame, "医疗费（元，诊疗费+医药费+住院费）：", 0)
        self.hospital_days = self.create_entry(medical_frame, "住院天数：", 1)
        self.meal_subsidy = self.create_entry(medical_frame, "住院伙食补助费（元/天，默认100）：", 2)
        self.nutrition_fee = self.create_entry(medical_frame, "营养费（元）：", 3)
        self.traffic_fee = self.create_entry(medical_frame, "交通费（元）：", 4)
        self.accommodation_days = self.create_entry(medical_frame, "住宿天数：", 5)
        self.follow_up_treatment_fee = self.create_entry(medical_frame, "后续治疗费（元）：", 6)
        
        # 误工费框架
        work_frame = ttk.LabelFrame(scrollable_frame, text="💼 误工费", padding=6)
        work_frame.pack(fill="x", padx=10, pady=3)
        
        self.work_income_type = self.create_combobox(work_frame, "收入类型", 
                                                     ["固定收入", "无固定收入（能证明最近三年平均）", "无固定收入（不能证明，参照行业平均）"], 0)
        # 绑定收入类型变化事件，显示/隐藏相关字段
        self.work_income_type.bind("<<ComboboxSelected>>", self.on_income_type_changed)
        
        self.monthly_income = self.create_entry(work_frame, "月收入（元，固定收入时填写）：", 1)
        self.avg_daily_income = self.create_entry(work_frame, "日均收入（元，无固定收入能证明时填写）：", 2)
        
        # 行业选择下拉框（仅在选择"无固定收入（不能证明，参照行业平均）"时显示）
        self.industry_label = tk.Label(work_frame, text="行业类型：", font=("Microsoft YaHei", 8))
        self.industry_label.grid(row=3, column=0, sticky="w", padx=6, pady=2)
        self.industry_type = ttk.Combobox(work_frame, values=list(self.INDUSTRY_SALARIES.keys()), 
                                          width=39, state="readonly", font=("Microsoft YaHei", 8))
        self.industry_type.grid(row=3, column=1, padx=6, pady=2)
        self.industry_type.set("其他行业")  # 默认值
        # 初始隐藏行业选择
        self.industry_label.grid_remove()
        self.industry_type.grid_remove()
        
        self.work_loss_days = self.create_entry(work_frame, "误工天数：", 4)
        
        # 护理费框架
        nursing_frame = ttk.LabelFrame(scrollable_frame, text="👨‍⚕️ 护理费", padding=6)
        nursing_frame.pack(fill="x", padx=10, pady=3)
        
        self.nursing_type = self.create_combobox(nursing_frame, "护理人员类型：", 
                                                 ["有收入", "无收入或雇佣护工"], 0)
        # 绑定护理人员类型变化事件，显示/隐藏相关字段
        self.nursing_type.bind("<<ComboboxSelected>>", self.on_nursing_type_changed)
        
        self.nursing_income = self.create_entry(nursing_frame, "护理人员日均收入（元，有收入时填写）：", 1)
        self.nursing_days = self.create_entry(nursing_frame, "护理天数：", 2)
        self.nursing_count = self.create_entry(nursing_frame, "护理人数（默认1人）：", 3)
        
        # 初始状态：根据默认选择显示/隐藏
        self.on_nursing_type_changed()
        
        # 残疾相关框架
        self.disability_frame = ttk.LabelFrame(scrollable_frame, text="♿ 残疾赔偿", padding=6)
        self.disability_frame.pack(fill="x", padx=10, pady=3)
        
        # 伤残等级输入（支持多处伤残，用逗号或分号分隔，如：5级,8级 或 3级;5级;9级）
        disability_label = tk.Label(self.disability_frame, text="伤残等级：", font=("Microsoft YaHei", 8))
        disability_label.grid(row=0, column=0, sticky="w", padx=6, pady=2)
        self.disability_level = tk.Entry(self.disability_frame, width=40, font=("Microsoft YaHei", 8))
        self.disability_level.grid(row=0, column=1, padx=6, pady=2)
        self.disability_level.insert(0, "无")
        # 添加提示标签
        hint_label = tk.Label(self.disability_frame, 
                             text="提示：支持多处伤残，用逗号或分号分隔，如：5级,8级 或 3级;5级;9级（最高等级在前）",
                             font=("Microsoft YaHei", 7), fg="#7f8c8d")
        hint_label.grid(row=0, column=2, padx=(3, 0), pady=2, sticky="w")
        
        self.disability_appliance_fee = self.create_entry(self.disability_frame, "残疾辅助器具费（元）：", 1)
        
        # 被扶养人生活费框架
        self.dependent_frame = ttk.LabelFrame(scrollable_frame, text="👨‍👩‍👧‍👦 被扶养人生活费", padding=6)
        self.dependent_frame.pack(fill="x", padx=10, pady=3)
        
        self.dependent_info = self.create_entry(self.dependent_frame, "被扶养人信息（格式：年龄1,扶养人数1;年龄2,扶养人数2，如：5,2;65,1）：", 0)
        tk.Label(self.dependent_frame, text="说明：不满18岁按(18-年龄)年计算；18-60岁无劳动能力按20年；60-75岁按[20-(年龄-60)]年；75岁以上按5年", 
                font=("Arial", 7), fg="gray").grid(row=1, column=0, columnspan=2, sticky="w", padx=6, pady=1)
        
        # 死亡相关框架
        death_frame = ttk.LabelFrame(scrollable_frame, text="⚰️ 死亡赔偿（如适用）", padding=6)
        death_frame.pack(fill="x", padx=10, pady=3)
        
        self.is_death = tk.BooleanVar()
        death_checkbutton = tk.Checkbutton(death_frame, text="是否死亡", variable=self.is_death,
                                           command=self.on_death_changed, font=("Microsoft YaHei", 8))
        death_checkbutton.grid(row=0, column=0, sticky="w", padx=6, pady=2)
        
        # 初始状态：如果死亡被选中，隐藏残疾赔偿
        self.on_death_changed()
        
        # 精神损害抚慰金框架
        mental_frame = ttk.LabelFrame(scrollable_frame, text="💔 精神损害抚慰金", padding=6)
        mental_frame.pack(fill="x", padx=10, pady=3)
        
        self.mental_damage = self.create_entry(mental_frame, "精神损害抚慰金（元）：", 0)
        
        # 按钮框架 - 紧凑设计
        button_container = tk.Frame(scrollable_frame, bg="#f8f9fa", relief="raised", bd=1)
        button_container.pack(fill="x", padx=10, pady=5)
        
        # 主操作按钮区域 - 横向排列
        main_button_frame = tk.Frame(button_container, bg="#f8f9fa")
        main_button_frame.pack(fill="x", padx=10, pady=5)
        
        # 计算赔偿按钮
        calculate_btn = tk.Button(main_button_frame, 
                                 text="✓ 计算赔偿", 
                                 command=self.calculate, 
                                 bg="#27ae60", fg="white", 
                                 font=("Microsoft YaHei", 11, "bold"),
                                 padx=20, pady=8, 
                                 relief="raised", bd=2,
                                 cursor="hand2", 
                                 activebackground="#229954",
                                 activeforeground="white",
                                 highlightthickness=0)
        calculate_btn.pack(side="left", padx=4, expand=True, fill="both")
        
        # 导出Word文档按钮
        export_btn = tk.Button(main_button_frame, 
                               text="📄 导出Word", 
                               command=self.export_to_word, 
                               bg="#3498db", fg="white", 
                               font=("Microsoft YaHei", 11, "bold"),
                               padx=20, pady=8, 
                               relief="raised", bd=2,
                               cursor="hand2", 
                               activebackground="#2980b9",
                               activeforeground="white",
                               highlightthickness=0)
        export_btn.pack(side="left", padx=4, expand=True, fill="both")
        
        # 清空数据按钮
        clear_btn = tk.Button(main_button_frame, 
                             text="🗑️ 清空", 
                             command=self.clear_all, 
                             bg="#95a5a6", fg="white", 
                             font=("Microsoft YaHei", 10, "bold"),
                             padx=15, pady=8, 
                             relief="raised", bd=2,
                             cursor="hand2", 
                             activebackground="#7f8c8d",
                             activeforeground="white",
                             highlightthickness=0)
        clear_btn.pack(side="left", padx=4, expand=True, fill="both")
        
        # 结果显示框架 - 紧凑设计
        result_frame = ttk.LabelFrame(scrollable_frame, text="📊 计算结果", padding=6)
        result_frame.pack(fill="both", expand=True, padx=10, pady=3)
        
        self.result_text = tk.Text(result_frame, height=10, wrap=tk.WORD, 
                                   font=("Consolas", 9), 
                                   bg="#ffffff", fg="#2c3e50",
                                   relief="solid", borderwidth=1)
        self.result_text.pack(fill="both", expand=True)
        
        # 存储计算结果和计算详情
        self.calculation_results = {}
        self.calculation_details = {}  # 存储详细的计算公式和步骤
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
    def create_entry(self, parent, label_text, row):
        """创建输入框"""
        label = tk.Label(parent, text=label_text, font=("Microsoft YaHei", 8))
        label.grid(row=row, column=0, sticky="w", padx=6, pady=2)
        entry = tk.Entry(parent, width=42, font=("Microsoft YaHei", 8),
                         relief="solid", borderwidth=1, bg="#ffffff")
        entry.grid(row=row, column=1, padx=6, pady=2)
        return entry
    
    def create_combobox(self, parent, label_text, values, row):
        """创建下拉框"""
        label = tk.Label(parent, text=label_text, font=("Microsoft YaHei", 8))
        label.grid(row=row, column=0, sticky="w", padx=6, pady=2)
        combobox = ttk.Combobox(parent, values=values, width=39, state="readonly",
                               font=("Microsoft YaHei", 8))
        combobox.grid(row=row, column=1, padx=6, pady=2)
        if values:
            combobox.set(values[0])
        return combobox
    
    def create_date_selectors(self, parent, label_text, row):
        """创建日期选择器（年、月、日三个下拉框）"""
        label = tk.Label(parent, text=label_text, font=("Microsoft YaHei", 8))
        label.grid(row=row, column=0, sticky="w", padx=6, pady=2)
        
        # 创建日期选择器框架
        date_frame = tk.Frame(parent)
        date_frame.grid(row=row, column=1, padx=6, pady=2, sticky="w")
        
        # 获取当前日期
        now = datetime.now()
        current_year = now.year
        current_month = now.month
        current_day = now.day
        
        # 生成年份列表（当前年份往前10年，往后2年）
        years = [str(y) for y in range(current_year - 10, current_year + 3)]
        # 生成月份列表
        months = [f"{m:02d}" for m in range(1, 13)]
        # 生成日期列表（默认31天，会根据月份动态调整）
        days = [f"{d:02d}" for d in range(1, 32)]
        
        # 创建年份下拉框
        year_label = tk.Label(date_frame, text="年", font=("Microsoft YaHei", 8))
        year_label.pack(side="left", padx=(0, 1))
        year_combo = ttk.Combobox(date_frame, values=years, width=6, 
                                 state="readonly", font=("Microsoft YaHei", 8))
        year_combo.set(str(current_year))
        year_combo.pack(side="left", padx=1)
        
        # 创建月份下拉框
        month_label = tk.Label(date_frame, text="月", font=("Microsoft YaHei", 8))
        month_label.pack(side="left", padx=(0, 1))
        month_combo = ttk.Combobox(date_frame, values=months, width=4, 
                                  state="readonly", font=("Microsoft YaHei", 8))
        month_combo.set(f"{current_month:02d}")
        month_combo.pack(side="left", padx=1)
        
        # 创建日期下拉框
        day_label = tk.Label(date_frame, text="日", font=("Microsoft YaHei", 8))
        day_label.pack(side="left", padx=(0, 1))
        day_combo = ttk.Combobox(date_frame, values=days, width=4, 
                                state="readonly", font=("Microsoft YaHei", 8))
        day_combo.set(f"{current_day:02d}")
        day_combo.pack(side="left", padx=1)
        
        # 更新日期列表的函数（根据年月调整天数）
        def update_days(*args):
            try:
                year = int(year_combo.get())
                month = int(month_combo.get())
                # 计算该月的天数
                if month in [1, 3, 5, 7, 8, 10, 12]:
                    max_day = 31
                elif month in [4, 6, 9, 11]:
                    max_day = 30
                else:  # 2月
                    # 判断闰年
                    if (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0):
                        max_day = 29
                    else:
                        max_day = 28
                
                # 更新日期列表
                new_days = [f"{d:02d}" for d in range(1, max_day + 1)]
                day_combo['values'] = new_days
                
                # 如果当前选择的日期超出范围，设置为该月最后一天
                current_day_val = day_combo.get()
                if not current_day_val or int(current_day_val) > max_day:
                    day_combo.set(f"{max_day:02d}")
            except:
                pass
        
        # 绑定年月变化事件
        year_combo.bind("<<ComboboxSelected>>", update_days)
        month_combo.bind("<<ComboboxSelected>>", update_days)
        
        # 初始化日期列表
        update_days()
        
        return year_combo, month_combo, day_combo
    
    def get_float_value(self, entry, default=0.0):
        """获取浮点数值"""
        try:
            value = entry.get().strip()
            return float(value) if value else default
        except ValueError:
            return default
    
    def get_int_value(self, entry, default=0):
        """获取整数值"""
        try:
            value = entry.get().strip()
            return int(value) if value else default
        except ValueError:
            return default
    
    def on_income_type_changed(self, event=None):
        """当收入类型改变时，显示/隐藏相关字段"""
        income_type = self.work_income_type.get()
        
        if income_type == "固定收入":
            # 显示月收入，隐藏日均收入和行业选择
            monthly_label = self.monthly_income.master.grid_slaves(row=1, column=0)
            if monthly_label:
                monthly_label[0].grid()
            self.monthly_income.grid()
            # 隐藏日均收入
            avg_label = self.avg_daily_income.master.grid_slaves(row=2, column=0)
            if avg_label:
                avg_label[0].grid_remove()
            self.avg_daily_income.grid_remove()
            # 隐藏行业选择
            if hasattr(self, 'industry_label'):
                self.industry_label.grid_remove()
            if hasattr(self, 'industry_type'):
                self.industry_type.grid_remove()
            
        elif income_type == "无固定收入（能证明最近三年平均）":
            # 显示日均收入，隐藏月收入和行业选择
            monthly_label = self.monthly_income.master.grid_slaves(row=1, column=0)
            if monthly_label:
                monthly_label[0].grid_remove()
            self.monthly_income.grid_remove()
            # 显示日均收入
            avg_label = self.avg_daily_income.master.grid_slaves(row=2, column=0)
            if avg_label:
                avg_label[0].grid()
            self.avg_daily_income.grid()
            # 隐藏行业选择
            if hasattr(self, 'industry_label'):
                self.industry_label.grid_remove()
            if hasattr(self, 'industry_type'):
                self.industry_type.grid_remove()
            
        else:  # 无固定收入（不能证明，参照行业平均）
            # 显示行业选择，隐藏月收入和日均收入
            monthly_label = self.monthly_income.master.grid_slaves(row=1, column=0)
            if monthly_label:
                monthly_label[0].grid_remove()
            self.monthly_income.grid_remove()
            # 隐藏日均收入
            avg_label = self.avg_daily_income.master.grid_slaves(row=2, column=0)
            if avg_label:
                avg_label[0].grid_remove()
            self.avg_daily_income.grid_remove()
            # 显示行业选择
            if hasattr(self, 'industry_label'):
                self.industry_label.grid()
            if hasattr(self, 'industry_type'):
                self.industry_type.grid()
    
    def on_nursing_type_changed(self, event=None):
        """当护理人员类型改变时，显示/隐藏相关字段"""
        nursing_type = self.nursing_type.get()
        
        if nursing_type == "有收入":
            # 显示护理人员日均收入输入框
            nursing_income_label = self.nursing_income.master.grid_slaves(row=1, column=0)
            if nursing_income_label:
                nursing_income_label[0].grid()
            self.nursing_income.grid()
        else:  # 无收入或雇佣护工
            # 隐藏护理人员日均收入输入框
            nursing_income_label = self.nursing_income.master.grid_slaves(row=1, column=0)
            if nursing_income_label:
                nursing_income_label[0].grid_remove()
            self.nursing_income.grid_remove()
    
    def calculate_multi_disability_coefficient(self, disability_levels_str):
        """
        计算多处伤残的伤残系数
        根据《道路交通事故受伤人员伤残评定》标准：
        1. 最高伤残等级系数：取所有伤残等级中最高的一个
        2. 附加指数：
           - 2-5级伤残：每处附加指数为4%
           - 6-10级伤残：每处附加指数为2%
           - 附加指数总和不超过10%
        3. 最终伤残系数 = 最高伤残等级系数 + 附加指数（但不超过100%）
        
        参数：
        - disability_levels_str: 伤残等级字符串，如"5级,8级"或"3级;5级;9级"
        
        返回：(最终伤残系数, 最高伤残等级, 附加指数, 计算详情)
        """
        if not disability_levels_str or disability_levels_str.strip() == "无":
            return 1.0, None, 0.0, "无伤残，系数为1.0"
        
        # 解析伤残等级
        disability_levels = []
        try:
            # 支持逗号和分号分隔
            parts = disability_levels_str.replace('，', ',').replace('；', ';').replace(',', ';').split(';')
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                # 提取数字
                if '级' in part:
                    level = int(part.replace('级', '').strip())
                else:
                    level = int(part.strip())
                if 1 <= level <= 10:
                    disability_levels.append(level)
        except (ValueError, AttributeError):
            return 1.0, None, 0.0, "伤残等级格式错误，按无伤残处理"
        
        if not disability_levels:
            return 1.0, None, 0.0, "无有效伤残等级，系数为1.0"
        
        # 去重并排序（从高到低）
        disability_levels = sorted(set(disability_levels))
        
        # 获取最高伤残等级
        max_level = disability_levels[0]  # 最高等级（数字最小）
        max_coefficient = self.DISABILITY_COEFFICIENTS.get(max_level, 1.0)
        
        # 计算附加指数（排除最高等级）
        additional_levels = disability_levels[1:] if len(disability_levels) > 1 else []
        additional_index = 0.0
        
        detail_parts = [f"伤残等级：{', '.join([f'{l}级' for l in disability_levels])}\n"]
        detail_parts.append(f"最高伤残等级：{max_level}级，系数：{max_coefficient}\n")
        
        if additional_levels:
            detail_parts.append("附加伤残等级：")
            for idx, level in enumerate(additional_levels):
                if idx > 0:
                    detail_parts.append("、")
                if 2 <= level <= 5:
                    additional_index += 0.04  # 4%
                    detail_parts.append(f"{level}级（附加4%）")
                elif 6 <= level <= 10:
                    additional_index += 0.02  # 2%
                    detail_parts.append(f"{level}级（附加2%）")
                # 注意：1级伤残不应作为附加等级，因为1级已经是100%
            
            # 附加指数总和不超过10%
            additional_index = min(additional_index, 0.10)
            detail_parts.append(f"\n附加指数合计：{additional_index * 100:.0f}%\n")
        else:
            detail_parts.append("无附加伤残等级\n")
        
        # 计算最终系数（不超过100%）
        final_coefficient = min(max_coefficient + additional_index, 1.0)
        
        detail_parts.append(f"最终伤残系数 = {max_coefficient} + {additional_index} = {final_coefficient}")
        if final_coefficient >= 1.0:
            detail_parts.append("（已达到100%上限）")
        
        detail = "".join(detail_parts)
        
        return final_coefficient, max_level, additional_index, detail
    
    def on_death_changed(self):
        """当死亡复选框状态改变时，显示/隐藏残疾赔偿框架"""
        if self.is_death.get():
            # 如果选择了死亡赔偿，隐藏残疾赔偿框架
            self.disability_frame.pack_forget()
            # 同时重置残疾赔偿相关字段
            if hasattr(self, 'disability_level'):
                if isinstance(self.disability_level, tk.Entry):
                    self.disability_level.delete(0, tk.END)
                    self.disability_level.insert(0, "无")
                else:
                    self.disability_level.set("无")
            if hasattr(self, 'disability_appliance_fee'):
                self.disability_appliance_fee.delete(0, tk.END)
        else:
            # 如果取消选择死亡赔偿，显示残疾赔偿框架
            # 在护理费框架之后、被扶养人生活费框架之前显示
            self.disability_frame.pack(fill="x", padx=15, pady=8, before=self.dependent_frame)
    
    def calculate_compensation_years(self, age):
        """
        计算赔偿年限
        根据年龄计算：60周岁以下按20年；60周岁以上每增加一岁减少一年；75周岁以上按5年
        """
        if age < 60:
            return 20
        elif age >= 75:
            return 5
        else:
            return 20 - (age - 60)
    
    def calculate_work_loss_fee(self):
        """
        计算误工费
        根据《最高人民法院关于审理人身损害赔偿案件适用法律若干问题的解释》第20条
        返回：(金额, 计算详情)
        """
        work_loss_days = self.get_int_value(self.work_loss_days)
        if work_loss_days <= 0:
            return 0, "误工天数为0，不计算误工费"
        
        income_type = self.work_income_type.get()
        
        if income_type == "固定收入":
            # 受害人有固定收入的，误工费按照实际减少的收入计算
            monthly_income = self.get_float_value(self.monthly_income)
            if monthly_income > 0:
                daily_income = monthly_income / 30
                amount = daily_income * work_loss_days
                detail = f"固定收入计算：\n月收入：{monthly_income:,.2f}元\n日均收入 = 月收入 ÷ 30 = {monthly_income:,.2f} ÷ 30 = {daily_income:,.2f}元/天\n误工费 = 日均收入 × 误工天数 = {daily_income:,.2f} × {work_loss_days} = {amount:,.2f}元"
                return amount, detail
            else:
                return 0, "月收入为0，不计算误工费"
        
        elif income_type == "无固定收入（能证明最近三年平均）":
            # 能证明最近三年平均收入的
            avg_daily_income = self.get_float_value(self.avg_daily_income)
            if avg_daily_income > 0:
                amount = avg_daily_income * work_loss_days
                detail = f"无固定收入（能证明最近三年平均）计算：\n最近三年平均日均收入：{avg_daily_income:,.2f}元/天\n误工费 = 日均收入 × 误工天数 = {avg_daily_income:,.2f} × {work_loss_days} = {amount:,.2f}元"
                return amount, detail
            else:
                return 0, "日均收入为0，不计算误工费"
        
        else:  # 无固定收入（不能证明，参照行业平均）
            # 不能证明的，参照受诉法院所在地相同或者相近行业上一年度职工的平均工资计算
            # 根据用户选择的行业获取对应的平均工资
            selected_industry = self.industry_type.get() if hasattr(self, 'industry_type') else "其他行业"
            industry_avg_salary = self.INDUSTRY_SALARIES.get(selected_industry, self.INDUSTRY_SALARIES['其他行业'])
            daily_avg_salary = industry_avg_salary / 365
            amount = daily_avg_salary * work_loss_days
            detail = f"无固定收入（不能证明，参照行业平均）计算\n选择行业：{selected_industry}\n行业平均工资：{industry_avg_salary:,.2f}元/年\n日均工资 = 年工资 ÷ 365 = {industry_avg_salary:,.2f} ÷ 365 = {daily_avg_salary:,.2f}元/天\n误工费 = 日均工资 × 误工天数 = {daily_avg_salary:,.2f} × {work_loss_days} = {amount:,.2f}元"
            return amount, detail
    
    def calculate_nursing_fee(self):
        """
        计算护理费
        根据《最高人民法院关于审理人身损害赔偿案件适用法律若干问题的解释》第21条
        返回：(金额, 计算详情)
        """
        nursing_days = self.get_int_value(self.nursing_days)
        nursing_count = self.get_int_value(self.nursing_count, 1)
        
        if nursing_days <= 0:
            return 0, "护理天数为0，不计算护理费"
        
        nursing_type = self.nursing_type.get()
        
        if nursing_type == "有收入":
            # 护理人员有收入的，参照误工费的规定计算
            nursing_income = self.get_float_value(self.nursing_income)
            if nursing_income > 0:
                amount = nursing_income * nursing_days * nursing_count
                detail = f"护理人员有收入计算：\n护理人员日均收入：{nursing_income:,.2f}元/天\n护理天数：{nursing_days}天\n护理人数：{nursing_count}人\n护理费 = 日均收入 × 护理天数 × 护理人数 = {nursing_income:,.2f} × {nursing_days} × {nursing_count} = {amount:,.2f}元"
                return amount, detail
            else:
                return 0, "护理人员日均收入为0，不计算护理费"
        else:
            # 护理人员没有收入或者雇佣护工的，参照当地护工从事同等级别护理的劳务报酬标准计算
            nursing_fee_per_day = self.STANDARDS['daily_nursing_fee']
            amount = nursing_fee_per_day * nursing_days * nursing_count
            detail = f"无收入或雇佣护工计算：\n护工标准：{nursing_fee_per_day:,.2f}元/天\n护理天数：{nursing_days}天\n护理人数：{nursing_count}人\n护理费 = 护工标准 × 护理天数 × 护理人数 = {nursing_fee_per_day:,.2f} × {nursing_days} × {nursing_count} = {amount:,.2f}元"
            return amount, detail
    
    def calculate_dependent_living_expense(self, victim_age, disability_coefficient=1.0, is_death=False):
        """
        计算被扶养人生活费
        根据《最高人民法院关于审理人身损害赔偿案件适用法律若干问题的解释》第28条
        2025年标准统一使用城镇居民人均消费支出
        
        计算公式：
        1. 不满18周岁：生活费 = 消费支出 × (18-实际年龄)
        2. 18-60周岁（无劳动能力）：生活费 = 消费支出 × 20年
        3. 60-75周岁：生活费 = 消费支出 × [20-(实际年龄-60)]年
        4. 75周岁以上：生活费 = 消费支出 × 5年
        5. 有其他扶养人时：赔偿义务人承担的费用 = 生活费 ÷ 扶养人数
        6. 被扶养人有数人时：年赔偿总额 ≤ 消费支出
        7. 需要考虑伤残系数：最终金额 = 计算金额 × 伤残系数
        8. 受害人死亡的，无需乘以伤残系数（视为系数100%）
        
        参数：
        - victim_age: 受害人年龄
        - disability_coefficient: 伤残系数（默认1.0，即无伤残）
        - is_death: 是否死亡（默认False）
        
        返回：(金额, 计算详情)
        """
        dependent_info_str = self.dependent_info.get().strip()
        if not dependent_info_str:
            return 0, "未填写被扶养人信息，不计算被扶养人生活费"
        
        base_consumption = self.STANDARDS['consumption']  # 统一使用城镇居民标准
        consumption_type = "广西上一年度城镇居民人均消费支出"
        
        # 解析被扶养人信息：格式为"年龄1,扶养人数1;年龄2,扶养人数2"
        dependents = []
        try:
            for item in dependent_info_str.split(';'):
                item = item.strip()
                if not item:
                    continue
                if ',' in item:
                    parts = item.split(',')
                    age = int(parts[0].strip())
                    support_count = int(parts[1].strip()) if len(parts) > 1 else 1
                    dependents.append({'age': age, 'support_count': support_count})
                else:
                    # 如果没有逗号，只有年龄，默认扶养人数为1
                    age = int(item)
                    dependents.append({'age': age, 'support_count': 1})
        except ValueError:
            return 0, "被扶养人信息格式错误"
        
        if not dependents:
            return 0, "未填写被扶养人信息，不计算被扶养人生活费"
        
        # 计算每个被扶养人的生活费年限和年生活费
        dependent_expenses = []
        detail_parts = [f"{consumption_type}：{base_consumption:,.2f}元/年\n"]
        
        for idx, dep in enumerate(dependents):
            age = dep['age']
            support_count = dep['support_count']
            
            # 计算该被扶养人的赔偿年限
            if age < 18:
                years = 18 - age
                age_desc = f"不满18周岁，按(18-{age})年计算"
            elif age >= 18 and age < 60:
                years = 20
                age_desc = f"18-60周岁（无劳动能力），按20年计算"
            elif age >= 60 and age < 75:
                years = 20 - (age - 60)
                age_desc = f"60-75周岁，按[20-({age}-60)]={years}年计算"
            else:  # 75岁以上
                years = 5
                age_desc = f"75周岁以上，按5年计算"
            
            if years <= 0:
                continue
            
            # 计算该被扶养人的年生活费（需要除以扶养人数）
            annual_expense_per_dependent = base_consumption / support_count
            
            dependent_expenses.append({
                'age': age,
                'years': years,
                'support_count': support_count,
                'annual_expense': annual_expense_per_dependent
            })
            
            detail_parts.append(f"被扶养人{idx+1}：{age}岁，{age_desc}，扶养人数{support_count}人\n年生活费 = {base_consumption:,.2f} ÷ {support_count} = {annual_expense_per_dependent:,.2f}元/年\n")
        
        if not dependent_expenses:
            return 0, "被扶养人信息无效"
        
        # 计算总费用，考虑年赔偿总额限制
        max_years = max(exp['years'] for exp in dependent_expenses)
        
        total_expense = 0
        year_details = []
        for year in range(max_years):
            year_total = 0
            active_deps = []
            for exp in dependent_expenses:
                if year < exp['years']:
                    year_total += exp['annual_expense']
                    active_deps.append(f"{exp['age']}岁")
            
            # 年赔偿总额不能超过消费支出
            original_total = year_total
            year_total = min(year_total, base_consumption)
            total_expense += year_total
            
            if year_total > 0:
                if original_total > base_consumption:
                    year_details.append(f"第{year+1}年：{'+'.join(active_deps)}的年生活费合计{original_total:,.2f}元，超过{base_consumption:,.2f}元，按{base_consumption:,.2f}元计算")
                else:
                    year_details.append(f"第{year+1}年：{'+'.join(active_deps)}的年生活费合计{year_total:,.2f}元")
        
        # 生成总计公式
        year_amounts = []
        for year in range(max_years):
            year_total = 0
            for exp in dependent_expenses:
                if year < exp['years']:
                    year_total += exp['annual_expense']
            year_total = min(year_total, base_consumption)
            if year_total > 0:
                year_amounts.append(f"{year_total:,.2f}")
        
        total_formula = " + ".join(year_amounts) if year_amounts else "0"
        
        # 应用伤残系数（死亡情况下系数为100%）
        original_total = total_expense
        total_expense = total_expense * disability_coefficient
        
        # 更新计算详情，说明考虑了伤残系数或死亡情况
        if is_death:
            # 死亡情况：系数为100%
            detail = "".join(detail_parts) + "\n按年计算明细：\n" + "\n".join(year_details) + f"\n\n小计 = " + total_formula + f" = {original_total:,.2f}元\n受害人死亡，系数为100%（无需乘以伤残系数）\n被扶养人生活费 = 小计 × 100% = {original_total:,.2f} × 1.0 = {total_expense:,.2f}元"
        elif disability_coefficient < 1.0:
            # 有伤残情况：使用伤残系数
            detail = "".join(detail_parts) + "\n按年计算明细：\n" + "\n".join(year_details) + f"\n\n小计 = " + total_formula + f" = {original_total:,.2f}元\n伤残系数：{disability_coefficient}\n被扶养人生活费 = 小计 × 伤残系数 = {original_total:,.2f} × {disability_coefficient} = {total_expense:,.2f}元"
        else:
            # 无伤残情况
            detail = "".join(detail_parts) + "\n按年计算明细：\n" + "\n".join(year_details) + f"\n\n总计 = " + total_formula + f" = {total_expense:,.2f}元"
        
        return total_expense, detail
    
    def calculate(self):
        """计算各项赔偿"""
        try:
            results = {}
            
            # 基本信息
            victim_name = self.victim_name.get().strip() or "未填写"
            victim_age = self.get_int_value(self.victim_age, 0)
            
            # 1. 医疗费 = 诊疗费+医药费+住院费
            medical_expense = self.get_float_value(self.medical_expense)
            results['医疗费'] = medical_expense
            if medical_expense > 0:
                self.calculation_details['医疗费'] = f"医疗费 = 诊疗费 + 医药费 + 住院费 = {medical_expense:,.2f}元"
            
            # 2. 后续治疗费
            follow_up_treatment_fee = self.get_float_value(self.follow_up_treatment_fee)
            results['后续治疗费'] = follow_up_treatment_fee
            if follow_up_treatment_fee > 0:
                self.calculation_details['后续治疗费'] = f"后续治疗费 = {follow_up_treatment_fee:,.2f}元"
            
            # 3. 住院伙食补助费
            hospital_days = self.get_int_value(self.hospital_days)
            meal_subsidy_per_day = self.get_float_value(self.meal_subsidy, 
                                                       self.STANDARDS['daily_meal_subsidy'])
            meal_subsidy_total = hospital_days * meal_subsidy_per_day
            results['住院伙食补助费'] = meal_subsidy_total
            if meal_subsidy_total > 0:
                self.calculation_details['住院伙食补助费'] = f"住院天数：{hospital_days}天\n补助标准：{meal_subsidy_per_day:,.2f}元/天\n住院伙食补助费 = 住院天数 × 补助标准 = {hospital_days} × {meal_subsidy_per_day:,.2f} = {meal_subsidy_total:,.2f}元"
            
            # 3. 营养费
            nutrition_fee = self.get_float_value(self.nutrition_fee)
            results['营养费'] = nutrition_fee
            if nutrition_fee > 0:
                self.calculation_details['营养费'] = f"营养费 = {nutrition_fee:,.2f}元"
            
            # 4. 交通费
            traffic_fee = self.get_float_value(self.traffic_fee)
            results['交通费'] = traffic_fee
            if traffic_fee > 0:
                self.calculation_details['交通费'] = f"交通费 = {traffic_fee:,.2f}元"
            
            # 5. 住宿费（330元/天 × 住宿天数）
            accommodation_days = self.get_int_value(self.accommodation_days)
            accommodation_fee_per_day = self.STANDARDS['daily_accommodation_fee']
            accommodation_fee = accommodation_days * accommodation_fee_per_day
            results['住宿费'] = accommodation_fee
            if accommodation_fee > 0:
                self.calculation_details['住宿费'] = f"住宿天数：{accommodation_days}天\n住宿费标准：{accommodation_fee_per_day:,.2f}元/天\n住宿费 = 住宿天数 × 住宿费标准 = {accommodation_days} × {accommodation_fee_per_day:,.2f} = {accommodation_fee:,.2f}元"
            
            # 6. 误工费（根据收入类型计算）
            work_loss_fee, work_detail = self.calculate_work_loss_fee()
            results['误工费'] = work_loss_fee
            self.calculation_details['误工费'] = work_detail
            
            # 7. 护理费（根据护理人员类型计算）
            nursing_fee_total, nursing_detail = self.calculate_nursing_fee()
            results['护理费'] = nursing_fee_total
            self.calculation_details['护理费'] = nursing_detail
            
            # 8. 残疾赔偿金（2025年标准统一使用城镇居民人均可支配收入，支持多处伤残）
            if isinstance(self.disability_level, tk.Entry):
                disability_level_str = self.disability_level.get().strip()
            else:
                disability_level_str = self.disability_level.get() if hasattr(self.disability_level, 'get') else "无"
            
            # 计算多处伤残系数
            disability_coefficient, max_level, additional_index, disability_detail = \
                self.calculate_multi_disability_coefficient(disability_level_str)
            
            if disability_coefficient < 1.0 or (disability_level_str and disability_level_str != "无"):
                base_income = self.STANDARDS['disposable_income']  # 统一使用城镇居民标准
                income_type = "广西上一年度城镇居民人均可支配收入"
                # 计算年限：根据年龄调整
                years = self.calculate_compensation_years(victim_age)
                disability_compensation = base_income * years * disability_coefficient
                results['残疾赔偿金'] = disability_compensation
                year_desc = f"{years}年" if victim_age < 60 else (f"{years}年（60周岁以上每增加一岁减少一年）" if victim_age < 75 else f"{years}年（75周岁以上按5年计算）")
                
                # 构建计算详情
                detail = f"{disability_detail}\n{income_type}：{base_income:,.2f}元/年\n赔偿年限：{year_desc}\n残疾赔偿金 = {income_type} × 赔偿年限 × 伤残系数 = {base_income:,.2f} × {years} × {disability_coefficient} = {disability_compensation:,.2f}元"
                self.calculation_details['残疾赔偿金'] = detail
            else:
                results['残疾赔偿金'] = 0
            
            # 9. 残疾辅助器具费
            disability_appliance_fee = self.get_float_value(self.disability_appliance_fee)
            results['残疾辅助器具费'] = disability_appliance_fee
            if disability_appliance_fee > 0:
                self.calculation_details['残疾辅助器具费'] = f"残疾辅助器具费 = {disability_appliance_fee:,.2f}元"
            
            # 10. 被扶养人生活费（按年龄段精确计算，2025年标准统一使用城镇居民人均消费支出，需要考虑伤残系数）
            # 注意：受害人死亡的，无需乘以伤残系数（视为系数100%）
            is_death = self.is_death.get()
            if is_death:
                # 死亡情况下，使用系数1.0（100%）
                dependent_coefficient = 1.0
            else:
                # 非死亡情况，使用伤残系数
                dependent_coefficient = disability_coefficient
            
            dependent_living_expense, dependent_detail = self.calculate_dependent_living_expense(victim_age, dependent_coefficient, is_death)
            results['被扶养人生活费'] = dependent_living_expense
            if dependent_living_expense > 0:
                self.calculation_details['被扶养人生活费'] = dependent_detail
            
            # 11. 死亡赔偿金（2025年标准统一使用城镇居民人均可支配收入）
            if self.is_death.get():
                base_income = self.STANDARDS['disposable_income']  # 统一使用城镇居民标准
                income_type = "广西上一年度城镇居民人均可支配收入"
                # 计算年限：根据年龄调整（60岁以上每增加一岁减少一年，75岁以上按5年）
                years = self.calculate_compensation_years(victim_age)
                death_compensation = base_income * years
                results['死亡赔偿金'] = death_compensation
                results['丧葬费'] = self.STANDARDS['funeral_expense']
                year_desc = f"{years}年" if victim_age < 60 else (f"{years}年（60周岁以上每增加一岁减少一年）" if victim_age < 75 else f"{years}年（75周岁以上按5年计算）")
                self.calculation_details['死亡赔偿金'] = f"{income_type}：{base_income:,.2f}元/年\n赔偿年限：{year_desc}\n死亡赔偿金 = {income_type} × 赔偿年限 = {base_income:,.2f} × {years} = {death_compensation:,.2f}元"
                self.calculation_details['丧葬费'] = f"丧葬费 = {self.STANDARDS['funeral_expense']:,.2f}元"
            else:
                results['死亡赔偿金'] = 0
                results['丧葬费'] = 0
            
            # 12. 精神损害抚慰金
            mental_damage = self.get_float_value(self.mental_damage)
            results['精神损害抚慰金'] = mental_damage
            if mental_damage > 0:
                self.calculation_details['精神损害抚慰金'] = f"精神损害抚慰金 = {mental_damage:,.2f}元"
            
            # 计算总计
            total = sum(results.values())
            results['总计'] = total
            
            # 生成总计的计算公式
            valid_items = [item for item in ['医疗费', '后续治疗费', '误工费', '护理费', '交通费', '住宿费', '住院伙食补助费', 
                          '营养费', '残疾赔偿金', '残疾辅助器具费', '被扶养人生活费', 
                          '死亡赔偿金', '丧葬费', '精神损害抚慰金']
                          if item in results and results[item] > 0]
            total_formula = " + ".join([f"{results[item]:,.2f}" for item in valid_items])
            self.calculation_details['总计'] = f"总计 = {total_formula} = {total:,.2f}元"
            
            # 保存结果
            self.calculation_results = results
            
            # 显示结果
            self.display_results(results, victim_name, victim_age)
            
            messagebox.showinfo("成功", "计算完成！请查看计算结果。")
            
        except Exception as e:
            messagebox.showerror("错误", f"计算过程中出现错误：{str(e)}")
            import traceback
            traceback.print_exc()
    
    def display_results(self, results, name, age):
        """显示计算结果"""
        self.result_text.delete(1.0, tk.END)
        
        output = f"{'='*50}\n"
        output += f"广西人身损害赔偿计算结果\n"
        output += f"{'='*50}\n\n"
        output += f"受害人姓名：{name}\n"
        output += f"受害人年龄：{age}岁\n"
        output += f"计算日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        output += f"\n{'-'*50}\n"
        output += f"各项赔偿明细：\n"
        output += f"{'-'*50}\n\n"
        
        # 按顺序显示各项赔偿
        items_order = ['医疗费', '后续治疗费', '误工费', '护理费', '交通费', '住宿费', '住院伙食补助费', 
                      '营养费', '残疾赔偿金', '残疾辅助器具费', '被扶养人生活费', 
                      '死亡赔偿金', '丧葬费', '精神损害抚慰金']
        
        for item in items_order:
            if item in results and results[item] > 0:
                output += f"{item:20s}：{results[item]:>15,.2f} 元\n"
        
        output += f"\n{'-'*50}\n"
        output += f"{'总计':20s}：{results['总计']:>15,.2f} 元\n"
        output += f"{'='*50}\n"
        
        self.result_text.insert(1.0, output)
    
    def export_to_word(self):
        """导出到Word文档"""
        if not self.calculation_results:
            messagebox.showwarning("警告", "请先进行计算！")
            return
        
        try:
            # 获取受害人姓名
            victim_name = self.victim_name.get().strip() or "未填写"
            # 如果姓名为"未填写"，使用默认名称
            if victim_name == "未填写":
                name_part = ""
            else:
                name_part = victim_name
            
            # 获取事故发生日期作为计算日期
            try:
                year = self.accident_date_year.get().strip()
                month = self.accident_date_month.get().strip()
                day = self.accident_date_day.get().strip()
                if year and month and day:
                    # 使用事故发生日期
                    date_part = f"{year}{month}{day}"
                else:
                    # 如果没有填写日期，使用当前日期
                    date_part = datetime.now().strftime('%Y%m%d')
            except:
                # 如果获取日期失败，使用当前日期
                date_part = datetime.now().strftime('%Y%m%d')
            
            # 生成文件名：受害人姓名+赔偿计算结果+计算日期
            if name_part:
                initial_filename = f"{name_part}赔偿计算结果{date_part}.docx"
            else:
                initial_filename = f"赔偿计算结果{date_part}.docx"
            
            # 选择保存位置
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")],
                initialfile=initial_filename
            )
            
            if not filename:
                return
            
            # 创建Word文档
            doc = Document()
            
            # 设置文档默认样式
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # 设置标题样式
            heading1 = doc.styles['Heading 1']
            heading1_font = heading1.font
            heading1_font.name = '黑体'
            heading1_font.size = Pt(16)
            heading1_font.bold = True
            heading1_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            heading2 = doc.styles['Heading 2']
            heading2_font = heading2.font
            heading2_font.name = '黑体'
            heading2_font.size = Pt(14)
            heading2_font.bold = True
            heading2_font._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            # 添加页脚（页码）- 使用标准方法
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 创建包含页码字段的段落
            p = footer_para._element
            p_r = OxmlElement('w:pPr')
            p.append(p_r)
            
            # 创建run
            r = OxmlElement('w:r')
            p.append(r)
            
            # 添加"第"字
            t = OxmlElement('w:t')
            t.text = '第 '
            r.append(t)
            
            # 添加页码字段开始标记
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            r.append(fldChar1)
            
            # 添加字段指令
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'
            r.append(instrText)
            
            # 添加页码字段结束标记
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            r.append(fldChar2)
            
            # 添加"页"字
            r2 = OxmlElement('w:r')
            p.append(r2)
            t2 = OxmlElement('w:t')
            t2.text = ' 页'
            r2.append(t2)
            
            # 设置字体
            for r_elem in p.findall(qn('w:r')):
                rPr = OxmlElement('w:rPr')
                r_elem.insert(0, rPr)
                font = OxmlElement('w:rFonts')
                font.set(qn('w:ascii'), '宋体')
                font.set(qn('w:eastAsia'), '宋体')
                font.set(qn('w:hAnsi'), '宋体')
                rPr.append(font)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')  # 10pt = 20 half-points
                rPr.append(sz)
            
            # 标题
            title = doc.add_heading('广西人身损害赔偿计算结果', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.name = '黑体'
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            
            doc.add_paragraph()  # 空行
            
            # 基本信息表格
            doc.add_heading('一、基本信息', level=1)
            victim_name = self.victim_name.get().strip() or "未填写"
            victim_age = self.get_int_value(self.victim_age, 0)
            # 获取日期（从三个下拉框获取）
            try:
                year = self.accident_date_year.get().strip()
                month = self.accident_date_month.get().strip()
                day = self.accident_date_day.get().strip()
                if year and month and day:
                    accident_date = f"{year}-{month}-{day}"
                else:
                    accident_date = "未填写"
            except:
                accident_date = "未填写"
            
            basic_table = doc.add_table(rows=3, cols=2)
            basic_table.style = 'Light Grid Accent 1'
            
            # 设置表格列宽
            basic_table.columns[0].width = Inches(2.0)
            basic_table.columns[1].width = Inches(4.5)
            
            basic_info = [
                ('受害人姓名', victim_name),
                ('受害人年龄', f"{victim_age}岁"),
                ('事故发生日期', accident_date),
            ]
            
            for i, (label, value) in enumerate(basic_info):
                # 设置标签单元格
                label_cell = basic_table.rows[i].cells[0]
                label_cell.text = label
                label_para = label_cell.paragraphs[0]
                label_para.runs[0].bold = True
                label_para.runs[0].font.name = '宋体'
                label_para.runs[0].font.size = Pt(12)
                label_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 设置值单元格
                value_cell = basic_table.rows[i].cells[1]
                value_cell.text = value
                value_para = value_cell.paragraphs[0]
                value_para.runs[0].font.name = '宋体'
                value_para.runs[0].font.size = Pt(12)
                value_para.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            doc.add_paragraph()  # 空行
            
            # 赔偿明细表格
            doc.add_heading('二、赔偿明细及计算公式', level=1)
            
            # 按顺序显示各项赔偿
            items_order = ['医疗费', '后续治疗费', '误工费', '护理费', '交通费', '住宿费', '住院伙食补助费', 
                          '营养费', '残疾赔偿金', '残疾辅助器具费', '被扶养人生活费', 
                          '死亡赔偿金', '丧葬费', '精神损害抚慰金']
            
            # 计算有效项目
            valid_items = [item for item in items_order 
                          if item in self.calculation_results and self.calculation_results[item] > 0]
            
            if valid_items:
                # 创建赔偿明细表格：序号、项目、金额、计算方式
                # 注意：rows需要包含表头，所以是len(valid_items) + 1
                detail_table = doc.add_table(rows=len(valid_items) + 1, cols=4)
                detail_table.style = 'Light Grid Accent 1'
                
                # 设置表格自动调整（根据内容）
                tbl = detail_table._tbl
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                
                # 设置表格宽度为100%（自动调整）
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '0')
                tblW.set(qn('w:type'), 'auto')
                tblPr.append(tblW)
                
                # 设置表格布局为自动调整
                tblLayout = OxmlElement('w:tblLayout')
                tblLayout.set(qn('w:type'), 'autofit')
                tblPr.append(tblLayout)
                
                # 设置表格列宽（前三列紧凑，最后一列宽松）
                detail_table.columns[0].width = Inches(0.4)  # 序号列（最紧凑）
                detail_table.columns[1].width = Inches(1.0)  # 项目列（紧凑，项目名称一般较短）
                detail_table.columns[2].width = Inches(1.0)  # 金额列（紧凑，金额格式固定）
                detail_table.columns[3].width = Inches(5.1)  # 计算方式列（宽松，容纳详细公式）
                
                # 设置单元格内边距（前三列紧凑，最后一列稍宽松）
                for row_idx, row in enumerate(detail_table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        tcPr = cell._element.tcPr
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            cell._element.insert(0, tcPr)
                        
                        # 前三列（序号、项目、金额）使用更小的内边距，最后一列（计算方式）使用稍大的内边距
                        if col_idx < 3:
                            # 前三列：紧凑内边距
                            left_right = '80'  # 约4pt
                            top_bottom = '50'  # 约2.5pt
                        else:
                            # 最后一列：稍宽松内边距
                            left_right = '120'  # 约6pt
                            top_bottom = '60'   # 约3pt
                        
                        tcMar = OxmlElement('w:tcMar')
                        for margin_name, margin_value in [('top', top_bottom), ('left', left_right), 
                                                          ('bottom', top_bottom), ('right', left_right)]:
                            margin = OxmlElement(f'w:{margin_name}')
                            margin.set(qn('w:w'), margin_value)
                            margin.set(qn('w:type'), 'dxa')
                            tcMar.append(margin)
                        tcPr.append(tcMar)
                
                # 表头
                header_cells = detail_table.rows[0].cells
                header_texts = ['序号', '项目', '金额', '计算方式']
                
                # 设置表头格式（加粗、居中、统一字体）
                for idx, cell in enumerate(header_cells):
                    # 先保存文本
                    header_text = header_texts[idx]
                    
                    # 清空单元格内容
                    cell.paragraphs[0].clear()
                    
                    # 重新添加文本
                    para = cell.paragraphs[0]
                    run = para.add_run(header_text)
                    run.bold = True
                    run.font.name = '黑体'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置单元格背景色（浅灰色）
                    tcPr = cell._element.get_or_add_tcPr()
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'E7E6E6')
                    shading_elm.set(qn('w:val'), 'clear')
                    tcPr.append(shading_elm)
                
                # 填充数据（注意：行索引从1开始，因为0是表头）
                for idx, item in enumerate(valid_items):
                    row = detail_table.rows[idx + 1]  # 从第2行开始（索引1）
                    
                    # 序号列
                    cell0 = row.cells[0]
                    cell0.paragraphs[0].clear()
                    para0 = cell0.paragraphs[0]
                    run0 = para0.add_run(str(idx + 1))
                    run0.font.name = '宋体'
                    run0.font.size = Pt(10)
                    run0._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 项目名称列
                    cell1 = row.cells[1]
                    cell1.paragraphs[0].clear()
                    para1 = cell1.paragraphs[0]
                    run1 = para1.add_run(item)
                    run1.font.name = '宋体'
                    run1.font.size = Pt(10)
                    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 金额列（右对齐，千分位格式）
                    cell2 = row.cells[2]
                    cell2.paragraphs[0].clear()
                    para2 = cell2.paragraphs[0]
                    amount_text = f"{self.calculation_results[item]:,.2f}"
                    run2 = para2.add_run(amount_text)
                    run2.font.name = '宋体'
                    run2.font.size = Pt(10)
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # 计算方式列（支持多行显示）
                    cell3 = row.cells[3]
                    cell3.paragraphs[0].clear()
                    para3 = cell3.paragraphs[0]
                    
                    if item in self.calculation_details:
                        detail = self.calculation_details[item]
                        # 如果包含换行符，分行显示
                        if '\n' in detail:
                            lines = detail.split('\n')
                            for i, line in enumerate(lines):
                                if i > 0:
                                    para3 = cell3.add_paragraph()
                                run3 = para3.add_run(line.strip())
                                run3.font.name = '宋体'
                                run3.font.size = Pt(9.5)
                                run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        else:
                            # 单行显示，将分号替换为换行
                            formula_text = detail.replace('；', '\n').replace(';', '\n')
                            if '\n' in formula_text:
                                lines = formula_text.split('\n')
                                for i, line in enumerate(lines):
                                    if i > 0:
                                        para3 = cell3.add_paragraph()
                                    run3 = para3.add_run(line.strip())
                                    run3.font.name = '宋体'
                                    run3.font.size = Pt(9.5)
                                    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            else:
                                run3 = para3.add_run(formula_text)
                                run3.font.name = '宋体'
                                run3.font.size = Pt(9.5)
                                run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        run3 = para3.add_run(f"{item} = {self.calculation_results[item]:,.2f} 元")
                        run3.font.name = '宋体'
                        run3.font.size = Pt(9.5)
                        run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    
                    para3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 设置行高（紧凑但不过于拥挤）
                    tr = row._element
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '300')  # 约15pt的行高
                    trHeight.set(qn('w:hRule'), 'atLeast')
                    trPr.append(trHeight)
                
                doc.add_paragraph()  # 空行
            
            # 总计表格
            doc.add_heading('三、赔偿总额', level=1)
            total_table = doc.add_table(rows=2, cols=2)
            total_table.style = 'Light Grid Accent 1'
            
            # 设置列宽
            total_table.columns[0].width = Inches(2.0)
            total_table.columns[1].width = Inches(5.0)
            
            # 表头
            total_table.rows[0].cells[0].text = '项目'
            total_table.rows[0].cells[1].text = '金额（元）'
            for cell in total_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 总计行
            total_table.rows[1].cells[0].text = '赔偿总额'
            total_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
            total_table.rows[1].cells[1].text = f"{self.calculation_results['总计']:,.2f}"
            total_table.rows[1].cells[1].paragraphs[0].runs[0].bold = True
            total_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 添加总计的计算公式
            if '总计' in self.calculation_details:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('计算公式：').bold = True
                doc.add_paragraph(self.calculation_details['总计'])
            
            # 计算依据
            doc.add_heading('四、计算依据', level=1)
            doc.add_paragraph('本计算依据以下法律法规及标准文件：')
            
            # 使用有序列表
            
            p1 = doc.add_paragraph('《广西壮族自治区道路交通事故损害赔偿项目及计算标准》（桂高法会〔2025〕13号）', style='List Number')
            p2 = doc.add_paragraph('《广西壮族自治区公安厅关于道路交通事故处理有关问题的通知》（桂公通〔2025〕60号）', style='List Number')
            
            doc.add_paragraph()
            doc.add_paragraph('注：2025年标准统一使用广西上一年度城镇居民人均可支配收入和城镇居民人均消费支出标准进行计算。')
            
            # 备注
            doc.add_heading('五、备注', level=1)
            doc.add_paragraph('1. 本计算结果仅供参考，实际赔偿金额以法院判决为准。')
            doc.add_paragraph('2. 各项费用需提供相应的票据和证明材料。')
            doc.add_paragraph('3. 误工费、护理费的计算方式已根据收入类型进行区分。')
            doc.add_paragraph('4. 被扶养人生活费的计算已考虑年赔偿总额限制。')
            doc.add_paragraph('5. 如对计算结果有疑问，请咨询广西瀛桂律师事务所唐学智律师，联系电话18078374299。')
            
            # 保存文档
            doc.save(filename)
            messagebox.showinfo("成功", f"Word文档已保存至：\n{filename}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出Word文档时出现错误：{str(e)}")
            import traceback
            traceback.print_exc()
    
    def clear_all(self):
        """清空所有数据"""
        if messagebox.askyesno("确认", "确定要清空所有数据吗？"):
            # 先重置日期选择器为当前日期
            try:
                now = datetime.now()
                self.accident_date_year.set(str(now.year))
                self.accident_date_month.set(f"{now.month:02d}")
                self.accident_date_day.set(f"{now.day:02d}")
            except:
                pass
            
            # 清空所有输入框
            for widget in self.root.winfo_children():
                self._clear_widget(widget)
            
            # 再次重置日期选择器（因为上面的清空可能会重置它）
            try:
                now = datetime.now()
                self.accident_date_year.set(str(now.year))
                self.accident_date_month.set(f"{now.month:02d}")
                self.accident_date_day.set(f"{now.day:02d}")
            except:
                pass
            
            # 确保死亡复选框取消后，残疾赔偿框架重新显示
            if hasattr(self, 'is_death') and not self.is_death.get():
                if hasattr(self, 'disability_frame'):
                    # 检查框架是否已隐藏
                    try:
                        self.disability_frame.pack_info()
                    except:
                        # 如果框架被隐藏了，重新显示
                        self.disability_frame.pack(fill="x", padx=15, pady=8, before=self.dependent_frame)
            
            self.result_text.delete(1.0, tk.END)
            self.calculation_results = {}
            self.calculation_details = {}
            messagebox.showinfo("提示", "数据已清空！")
    
    def _clear_widget(self, widget):
        """递归清空组件"""
        if isinstance(widget, tk.Entry):
            widget.delete(0, tk.END)
        elif isinstance(widget, ttk.Combobox):
            # 重置为第一个选项
            values = widget['values']
            if values:
                widget.set(values[0])
        elif isinstance(widget, tk.Checkbutton):
            widget.deselect()
        elif hasattr(widget, 'winfo_children'):
            for child in widget.winfo_children():
                self._clear_widget(child)


def main():
    """主函数"""
    root = tk.Tk()
    app = GuangxiCompensationCalculator(root)
    root.mainloop()


if __name__ == "__main__":
    main()
